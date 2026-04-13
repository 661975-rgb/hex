"""
AI 企業級投資戰情室 - 旗艦版 (Pro Plus)
整合：自動雷達、2.5 世代核心、兩年份深度數據回測、專業 Word/TXT 匯出。
"""
import os
import sys
import subprocess
import logging
import io
import time
from typing import Tuple, List, Dict, Any

# ==========================================
# 0. 魔法啟動器：全自動環境適應 (隨身碟相容)
# ==========================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

if "streamlit" not in sys.modules and __name__ == "__main__":
    req_file = os.path.join(BASE_DIR, "requirements.txt")
    if not os.path.exists(req_file):
        with open(req_file, "w", encoding="utf-8") as f:
            f.write("streamlit\nyfinance\nplotly\ngoogle-genai\npandas\ntenacity\npython-docx\n")
    try:
        import streamlit
        import yfinance
        import plotly
        import docx
    except ImportError:
        print("📦 偵測到新環境，正在安裝專業金融組件...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "-r", req_file, "--quiet"])
    subprocess.check_call([sys.executable, "-m", "streamlit", "run", os.path.abspath(__file__)])
    sys.exit(0)

# ==========================================
# 1. 核心套件載入
# ==========================================
import streamlit as st
import yfinance as yf
import pandas as pd
import plotly.graph_objects as go
from google import genai
from google.genai import types
from tenacity import retry, stop_after_attempt, wait_exponential
from docx import Document

# 設定日誌
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

st.set_page_config(page_title="AI 企業戰情室 Pro", page_icon="🏦", layout="wide")

# ==========================================
# 2. 🔐 身份驗證與狀態管理
# ==========================================
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    st.title("🔒 企業機密戰情室 - 身份驗證")
    pwd_input = st.text_input("🔑 請輸入主管通關密碼", type="password")
    if st.button("解鎖進入", type="primary"):
        if pwd_input == "boss888":
            st.session_state['logged_in'] = True
            st.rerun()
        else:
            st.error("❌ 密碼錯誤，拒絕訪問。")
    st.stop()

if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "assistant", "content": "總監您好！我是您的 AI 顧問。深度數據庫已就緒，請輸入股票代號開始分析。"}]

# ==========================================
# 3. 核心引擎 (API & 數據抓取)
# ==========================================
def get_api_key() -> str:
    try:
        if "GEMINI_API_KEY" in st.secrets: return st.secrets["GEMINI_API_KEY"]
    except: pass
    key_file = os.path.join(BASE_DIR, "KEY.TXT")
    if os.path.exists(key_file):
        with open(key_file, "r", encoding="utf-8") as f: return f.read().strip()
    st.error("❌ 找不到 KEY.TXT"); st.stop()

@st.cache_data(ttl=300)
def fetch_sidebar_info(ticker: str):
    """抓取側邊欄即時資訊"""
    try:
        stock = yf.Ticker(ticker)
        info = stock.info
        curr = info.get('currentPrice', info.get('regularMarketPrice', 'N/A'))
        prev = info.get('previousClose', curr)
        diff = curr - prev if isinstance(curr, (int, float)) else 0
        pct = (diff / prev) * 100 if prev != 0 else 0
        name = info.get('longName', info.get('shortName', ticker))
        return name, curr, diff, pct
    except: return ticker, "N/A", 0, 0

@st.cache_data(ttl=3600)
def fetch_main_data(ticker: str, period: str):
    """抓取主圖表數據"""
    df = yf.Ticker(ticker).history(period=period)
    if not df.empty:
        df.reset_index(inplace=True)
        if pd.api.types.is_datetime64tz_dtype(df['Date']):
            df['Date'] = df['Date'].dt.tz_localize(None)
    return df

def get_best_models(client: genai.Client) -> list[str]:
    """動態雷達優先權"""
    fallback = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-flash-latest"]
    try:
        models_gen = client.models.list()
        valid = [m.name.replace("models/", "") for m in models_gen]
        valid.sort(key=lambda n: 0 if "2.5" in n else (1 if "2.0" in n else 2))
        return valid[:3]
    except: return fallback

def ai_call(client, model, messages, instruction, max_t):
    """底層 API 封裝"""
    config = types.GenerateContentConfig(
        system_instruction=instruction,
        temperature=0.2,
        max_output_tokens=max_t,
        tools=[{"google_search": {}}]
    )
    return client.models.generate_content(model=model, contents=messages, config=config).text

# ==========================================
# 4. 🎛️ 左側專業控制台 (Sidebar)
# ==========================================
with st.sidebar:
    st.markdown("### 🏢 系統監控")
    st.success("🟢 節點：隨身碟加密連線中")
    st.info("🤖 核心：Gemini 2.5 (深度分析版)")
    st.divider()

    st.markdown("### 📈 即時行情")
    ticker = st.text_input("股票代號", "2330.TW").upper()
    name, price, diff, pct = fetch_sidebar_info(ticker)
    st.markdown(f"**{name}**")
    color = "red" if diff >= 0 else "green" # 台股視覺
    st.markdown(f"<h2 style='color:{color}; margin:0;'>{price}</h2>", unsafe_allow_html=True)
    st.markdown(f"<span style='color:{color};'>{'+' if diff >= 0 else ''}{diff:.2f} ({pct:.2f}%)</span>", unsafe_allow_html=True)
    
    st.divider()
    st.markdown("### ⚙️ 戰略配置")
    period = st.selectbox("圖表區間", ["1mo", "3mo", "6mo", "1y", "2y"], index=3)
    
    st.divider()
    st.markdown("### 📄 文檔中心")
    if 'report_text' in st.session_state:
        st.caption("✅ 深度報告已就緒")
        # Word 匯出
        doc = Document()
        doc.add_heading(f"{ticker} 企業級財務分析報告", 0)
        doc.add_paragraph(st.session_state['report_text'])
        bio = io.BytesIO(); doc.save(bio)
        st.download_button("📄 下載 Word 報告", bio.getvalue(), f"{ticker}_Analysis.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        # TXT 匯出
        txt_b = io.StringIO(); txt_b.write(st.session_state['report_text'])
        st.download_button("📝 下載純文字 TXT", txt_b.getvalue().encode('utf-8'), f"{ticker}_Report.txt", "text/plain", use_container_width=True)
    else:
        st.caption("⏳ 尚未生成專業報告")
    
    st.divider()
    if st.button("🚪 安全登出", use_container_width=True):
        st.session_state.logged_in = False; st.rerun()

# ==========================================
# 5. 右側主介面 (Main Area)
# ==========================================
df = fetch_main_data(ticker, period)
st.title(f"戰術終端：{ticker} 深度分析")

if not df.empty:
    fig = go.Figure(data=[go.Candlestick(x=df['Date'], open=df['Open'], high=df['High'], low=df['Low'], close=df['Close'], increasing_line_color='red', decreasing_line_color='green')])
    fig.update_layout(template='plotly_dark', margin=dict(l=0, r=0, t=0, b=0), height=350)
    st.plotly_chart(fig, use_container_width=True)

# ⚡ 生成報告按鈕
st.markdown("---")
if st.button("⚡ 啟動 AI 深度財務評估 (含 2025 數據回測)", type="primary", use_container_width=True):
    with st.spinner("正在抓取 2 年份深度數據並撰寫長篇報告..."):
        try:
            client = genai.Client(api_key=get_api_key())
            # 修正：專門為報告抓取 2 年數據
            deep_df = yf.Ticker(ticker).history(period="2y")
            deep_ctx = deep_df.tail(250)[['Close', 'Volume']].to_string() # 提供足夠的歷史資料
            
            report_prompt = f"""
            你是一位資深外資財務分析師。請針對 {ticker} 撰寫一份超過 1000 字的深度專業報告。
            你必須根據提供的 2 年份數據進行年度對比（特別是 2025 年與當前的漲跌表現）。
            
            請務必包含以下章節，且內容必須詳實、具備數據支撐：
            一、執行摘要 (Executive Summary)
            二、財報分析 (Financial Statement Analysis) - 包含獲利能力與趨勢
            三、技術面分析 (Technical Analysis) - 包含 2025 至現在的價格走勢
            四、基本面亮點 (Fundamental Highlights)
            五、投資風險評估 (Investment Risk Assessment)
            六、風險提示與最終建議 (Risk Warning)
            """
            
            model = get_best_models(client)[0]
            api_msg = [types.Content(role="user", parts=[types.Part.from_text(text=f"歷史數據庫：\n{deep_ctx}")])]
            # 解鎖 Token 至 4000
            resp = ai_call(client, model, api_msg, report_prompt, 4000)
            st.session_state['report_text'] = resp
            st.success("✅ 深度報告生成完畢！請前往左側面板下載。")
        except Exception as e:
            st.error(f"分析失敗: {e}")

if 'report_text' in st.session_state:
    with st.expander("👁️ 預覽深度報告內容 (點擊下載按鈕可匯出完整檔案)"):
        st.markdown(st.session_state['report_text'])

# 對話區 (節流)
st.divider()
st.subheader("🤖 投資策略諮詢")
for m in st.session_state.messages:
    with st.chat_message(m["role"]): st.markdown(m["content"])

if prompt := st.chat_input("詢問有關 2025 表現或未來展望..."):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"): st.markdown(prompt)
    with st.chat_message("assistant"):
        with st.spinner("思考中..."):
            client = genai.Client(api_key=get_api_key())
            sys_inst = f"你是分析師，針對 {ticker} 回答。請精簡。數據背景：{df.tail(10).to_string()}"
            # 對話維持 500 Token 節流
            history = st.session_state.messages[-4:]
            api_msgs = [types.Content(role=("user" if h["role"]=="user" else "model"), parts=[types.Part.from_text(text=h["content"])]) for h in history]
            model = get_best_models(client)[0]
            resp = ai_call(client, model, api_msgs, sys_inst, 500)
            st.markdown(resp)
            st.session_state.messages.append({"role": "assistant", "content": resp})